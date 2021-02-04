from datetime import datetime,timedelta
from dateutil.parser import parse
import Logic
from docx import Document
import os
from PyQt5.QtGui import QColor,QBrush
from PyQt5.QtWidgets import QMessageBox,QDialog

Head_label = ['装车地点', '作业线路', '装车去向', '配空车次', '配空车数', '实装重车', '调妥时间',
              '封堵开始', '封堵结束', '装车开始', '装车完毕', '平车开始', '平车结束', '具备挂车条件','挂车时间', '线内作业时间分析','待挂时间分析']

Today = datetime.today()
# 定义today 为“2020.xx.xx”格式的字符串
today =  Today.strftime('%Y.%m.%d')
tomorrow = Today + timedelta(days=1)
tomorrow = tomorrow.strftime('%Y.%m.%d')
#=======================================================================================================================



def table_analyse(table,table_time):
    df = Logic.Table.table_to_df(table,table_time)
    time_col_list = ['调妥时间', '封堵开始', '封堵结束', '装车开始', '装车完毕', '平车开始', '平车结束', '具备挂车条件', '挂车时间']
    for time_col_name in time_col_list:
        df[time_col_name] = df[time_col_name].fillna('')
        df[time_col_name] = df[time_col_name].astype('str')
        df[time_col_name] = '{} '.format(table_time) + df[time_col_name]
        df[time_col_name] = df[time_col_name].replace('{} '.format(table_time),'')
        for i, v in df[time_col_name].items():
            try:
                time_time = parse(v)
                # time_str = datetime.strftime(time_time,'%Y-%m-%d %H:%M')
                df.loc[i, time_col_name] = time_time
            except:
                pass
    df1 = df
    #==============================
    def time_lessthan_zero(sing_cell):
        if sing_cell < 0:
            sing_cell = sing_cell + 24
        return sing_cell

    def series_round(single_cell):
        return round(single_cell,2)

    for i in range(26):
        try:
            df1.loc[i, '线内用时'] = (df1.loc[i, '具备挂车条件'] - df1.loc[i, '调妥时间'])/timedelta(minutes=60)
            df1['线内用时'] = df1['线内用时'].apply(series_round)
            df1['线内用时'] = df1['线内用时'].apply(time_lessthan_zero)
        except TypeError:
            df1.loc[i, '线内用时'] = 0
    for i in range(26):
        try:
            time = (df1.loc[i, '挂车时间'] - df1.loc[i, '具备挂车条件'])/timedelta(minutes=60)
            df1.loc[i, '待挂用时'] = time
            print(time,'\n',i)
            df1['待挂用时'] = df1['待挂用时'].apply(series_round)
            df1['待挂用时'] = df1['待挂用时'].apply(time_lessthan_zero)
        except TypeError:
            df1.loc[i, '待挂用时'] = 0
    df1 = df1.iloc[:,1:]
    try:
        df1.to_excel('{}分析表.xlsx'.format(table_time))
    except PermissionError:
        remark_dialog = QDialog()
        QMessageBox.critical(remark_dialog, "注意", "请先关闭已打开的分析表", QMessageBox.Ok | QMessageBox.Cancel,
                             QMessageBox.Ok)

    return df1


def check_if_overtime(table_df,table_time):
    #检查是否作业超时
    table_df1 = table_analyse(table_df,table_time)
    try:
        condition_overtime = (table_df1['待挂用时'] > 2) | (table_df1['线内用时'] > 4.7)
        use_df = table_df1[condition_overtime]
        wait_pull_overtime_table = use_df[use_df['待挂用时'] > 2]
        work_overtime_table = use_df[use_df['线内用时'] > 4.7]

        work_overtime_table_dict = {index:time for index,time in work_overtime_table['线内用时'].items()}
        wait_pull_overtime_table_dict = {index:time for index,time in wait_pull_overtime_table['待挂用时'].items()}
        the_cell_with_background_color_index = []
        if len(work_overtime_table_dict) != 0:
            for k,v in work_overtime_table_dict.items():
                table_df.item(k,0).setBackground(QBrush(QColor(255, 0, 0)))
                the_cell_with_background_color_index.append(k)
        if len(wait_pull_overtime_table_dict) != 0:
            for k,v in wait_pull_overtime_table_dict.items():
                table_df.item(k,0).setBackground(QBrush(QColor(255,127,0)))
                the_cell_with_background_color_index.append(k)
        for index in range(0,26):
            if index not in the_cell_with_background_color_index:
                table_df.item(index,0).setBackground(QBrush(QColor(0x00,0xff,0x00,0x00)))
            else:
                pass
    except ImportError:
        pass

class Report_docx:
    def __init__(self,table,table_time):
        self.table = table
        self.table_time = table_time
        self.path = './表格'

    def use_table(self):
        use_df = table_analyse(self.table,self.table_time)
        return use_df

    def wash_use_table(self):
        # 生成仅包含超时列的excel
        use_df = self.use_table()
        condition_overtime = (use_df['待挂用时']>2)|(use_df['线内用时']>4.7)
        use_df = use_df[condition_overtime]
        wait_pull_overtime_table = use_df[use_df['待挂用时']>2]
        work_overtime_table = use_df[use_df['线内用时']>4.7]
        if not os.path.exists(self.path):
            os.makedirs('./表格')
        use_df.to_excel('./表格/haha.xlsx')


        document = Document()
        replace_dict = {'专一':'实业1线','专二':'实业2线','专三':'实业3线','专四':'实业4线',
                        '矿一':'矿三1线','矿二':'矿三2线'}

        document.add_paragraph('{}第_班矿三批车_列，装好_列，对妥_列。实业一期批车_列，装好_列，对妥_列。实业二期批_列，装好_列，对好_列。'.format(self.table_time))

        document.add_paragraph('一、{}第_班装车超时较多的原因分析'.format(self.table_time))
        if work_overtime_table.shape[0] == 0:
            document.add_paragraph('无')
        else:
            count = 0
            for row in work_overtime_table.index:
                work_time_hour = int(float(work_overtime_table.loc[row, '线内用时']) * 60 // 60)
                work_time_min = int(float(work_overtime_table.loc[row,'线内用时'])*60 - (work_time_hour*60))
                diaotuo_time = datetime.strftime(work_overtime_table.loc[row,'调妥时间'],'%H:%M')
                jubeiguachetiaojian_time = datetime.strftime(work_overtime_table.loc[row,'具备挂车条件'],'%H:%M')
                reason = work_overtime_table.loc[row,'线内作业时间分析']
                document.add_paragraph('{}、{}，{}{}次，{}调妥，{}具备挂车条件，共计用时{}小时{}分钟。'
                                               '\n原因：{}'.format(count +1,#条目数
                                                            replace_dict[work_overtime_table.loc[row,'作业线路']],#作业线路
                                                            work_overtime_table.loc[row, '装车去向'],#装车去向
                                                            work_overtime_table.loc[row,'配空车次'],#配空车次
                                                            diaotuo_time,#调妥时间
                                                            jubeiguachetiaojian_time,#具备挂车条件
                                                            work_time_hour,#线内作业的小时数
                                                            work_time_min,
                                                            reason))#线内作业的分钟数
                count += 1#给doc文档段落条目数计数使用

        document.add_paragraph('二、{}第_班待挂重车、待送空车时间较长原因分析'.format(self.table_time))
        if wait_pull_overtime_table.shape[0] == 0:
            document.add_paragraph('无')
        else:
            count = 0
            for row in wait_pull_overtime_table.index:
                wait_pull_hour = int(float(wait_pull_overtime_table.loc[row, '待挂用时']) * 60 // 60)
                wait_pull_min = int(float(wait_pull_overtime_table.loc[row,'待挂用时'])*60 - (wait_pull_hour*60))
                jubeiguachetiaojian_time = datetime.strftime(wait_pull_overtime_table.loc[row, '具备挂车条件'], '%H:%M')#具备挂车条件时间 ： str
                guache_time = datetime.strftime(wait_pull_overtime_table.loc[row, '挂车时间'], '%H:%M')#挂车时间 ：str
                reason = wait_pull_overtime_table.loc[row,'待挂时间分析']
                document.add_paragraph('{}、{}，{}{}次，{}具备挂车条件，{}挂出，重车待挂时间共计{}小时{}分钟。'
                                       '\n原因：{}'.format(count + 1,
                                                      replace_dict[wait_pull_overtime_table.loc[row, '作业线路']],
                                                      wait_pull_overtime_table.loc[row, '装车去向'],
                                                      wait_pull_overtime_table.loc[row, '配空车次'],
                                                      jubeiguachetiaojian_time,
                                                      guache_time,
                                                      wait_pull_hour,#待挂时间的小时数
                                                      wait_pull_min,
                                                      reason))#待挂时间的分钟数
                count += 1#给doc文档段落条目数计数使用
        document.add_paragraph('{}曹妃甸南第_班卸车_，其中煤_，钢材_。待卸_，其中_煤，_钢。'.format(self.table_time))
        try:
            document.save('./表格/{}写实.docx'.format(self.table_time))
        except PermissionError:
            remark_dialog = QDialog()
            QMessageBox.critical(remark_dialog, "注意", "请先关闭已打开的写实", QMessageBox.Ok | QMessageBox.Cancel,
                                 QMessageBox.Ok)
